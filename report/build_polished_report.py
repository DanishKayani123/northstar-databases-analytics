from __future__ import annotations

import csv
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/danishkayani/Documents/Codex/2026-05-23/files-mentioned-by-the-user-msg")
REPORT_DIR = ROOT / "report"
ASSETS_DIR = REPORT_DIR / "assets"
EVIDENCE_DIR = ASSETS_DIR / "evidence"
OUTPUTS_DIR = ROOT / "artifacts" / "outputs"
DATA_DIR = ROOT / "northstar_dataset"
OUTPUT_DOCX = REPORT_DIR / "northstar_assignment_report_polished.docx"


ACCENT = RGBColor(176, 83, 16)
TEXT = RGBColor(31, 41, 55)
MUTED = RGBColor(95, 99, 104)
LIGHT = RGBColor(245, 239, 230)
PALE = RGBColor(255, 247, 237)
BORDER = "D6D3D1"


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def set_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, font_size: float = 10.5) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = TEXT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, bold, color in [
        ("Title", 24, True, TEXT),
        ("Subtitle", 12, False, MUTED),
        ("Heading 1", 15, True, ACCENT),
        ("Heading 2", 12.5, True, TEXT),
        ("Heading 3", 11.5, True, TEXT),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NorthStar Databases and Analytics Coursework | Page ")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED
    set_page_number(p)


def add_cover_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("\n")

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Operational Diagnosis and Database Redesign for NorthStar Urban Mobility and Logistics")

    sub = doc.add_paragraph()
    sub.style = "Subtitle"
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Databases and Analytics Coursework Assignment")

    module_block = [
        "Student Name: [Add your name]",
        "Student ID: [Add your student ID]",
        "Module: CP60056E Databases and Analytics",
        "Assessment Component: Coursework Assignment",
        "Assessment Weighting: 80%",
        "Academic Year: 2025-2026",
        "Case Study: NorthStar Urban Mobility and Logistics",
        "GitHub Repository: [Add GitHub repo link after upload]",
        "Submission Route: Blackboard",
    ]
    for line in module_block:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(11)
        run.font.color.rgb = TEXT

    doc.add_paragraph()
    summary = doc.add_table(rows=1, cols=1)
    summary.alignment = WD_TABLE_ALIGNMENT.CENTER
    summary.autofit = False
    summary.columns[0].width = Inches(6.1)
    cell = summary.cell(0, 0)
    shade_cell(cell, "F5EFE6")
    set_cell_text(
        cell,
        "This report addresses the NorthStar case study through SQL in R, Python analytics, and a MongoDB-oriented redesign. "
        "The analysis focuses on delivery reliability, customer experience, data-quality failures, and the indexing strategy required "
        "to support a scalable operational reporting model.",
        font_size=11,
    )
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(
        "Written submission prepared to evidence structured querying, analytical interpretation, document modelling, and query optimisation."
    )
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.italic = True
    run.font.color.rgb = MUTED

    doc.add_page_break()


def add_front_matter(doc: Document) -> None:
    doc.add_paragraph("Table of Contents", style="Heading 1")
    entries = [
        "Abstract",
        "Chapter 1: Introduction and Assessment Context",
        "Chapter 2: Case Study Review, Data Inventory, and Methodology",
        "Chapter 3: SQL in R Analysis and Relational Findings",
        "Chapter 4: Python Analytics, Feature Engineering, and Interpretation",
        "Chapter 5: MongoDB Design, Indexing Strategy, and Query Optimisation",
        "Chapter 6: Recommendations and Conclusion",
        "References",
        "Appendix A: Evidence Figure Register",
        "Appendix B: Output Files Generated",
    ]
    for entry in entries:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.15)
        p.add_run(entry)

    doc.add_paragraph()
    doc.add_paragraph("List of Figures", style="Heading 1")
    figures = [
        "Figure 1. Python merged dataset preview used before feature engineering.",
        "Figure 2. SQL query result ranking hubs by non-on-time delivery share.",
        "Figure 3. SQL query result for orders with no linked delivery record.",
        "Figure 4. Zone-level performance table after cleaning inconsistent location labels.",
        "Figure 5. Failed delivery share by operating zone.",
        "Figure 6. Complaint mix by delivery status.",
        "Figure 7. Proof-of-completion risk comparison.",
        "Figure 8. Vehicle maintenance performance summary.",
        "Figure 9. Complaint intensity by hub.",
        "Figure 10. Highest app latency segments.",
        "Figure 11. Hub profitability proxy after direct cost and compensation.",
        "Figure 12. Proposed MongoDB customer case document.",
        "Figure 13. MongoDB index design evidence.",
        "Figure 14. MongoDB aggregation output for delayed or failed open cases.",
    ]
    for item in figures:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.15)
        run = p.add_run(item)
        run.font.size = Pt(10.5)

    doc.add_page_break()


def add_abstract(doc: Document) -> None:
    doc.add_paragraph("Abstract", style="Heading 1")
    doc.add_paragraph(
        "This report investigates the operational, customer-service, and data-management weaknesses described in the NorthStar Urban Mobility and Logistics case study. "
        "The assignment is approached as an integrated database and analytics problem rather than as a collection of isolated tables, because the central issue in the case "
        "is that business truth is fragmented across orders, deliveries, complaints, incidents, and app-event records. The analysis uses SQL in an R-oriented workflow to "
        "evaluate structured operational performance, Python to standardise inconsistent fields and engineer risk-sensitive features, and a MongoDB-oriented document model to "
        "restructure semi-structured customer cases and event histories. The main findings are that NorthStar has a severe order-to-delivery visibility gap, with 300 orders "
        "missing a delivery record, that the central operational layer is the weakest part of the service network, and that missing proof of completion and active use of vehicles "
        "already marked InRepair are both strong warning indicators of failure. The report concludes that NorthStar should retain relational processing for highly structured operational "
        "facts while redesigning complaint, exception, and event histories into document-based case records supported by targeted MongoDB indexes. This hybrid approach is the strongest "
        "fit for the dataset because it combines reliable reporting with more realistic representation of complex operational journeys."
    )
    doc.add_page_break()


def add_dataset_inventory_table(doc: Document) -> None:
    counts = {
        "orders": len(pd.read_csv(DATA_DIR / "orders.csv")),
        "deliveries": len(pd.read_csv(DATA_DIR / "deliveries.csv")),
        "customers": len(pd.read_csv(DATA_DIR / "customers.csv")),
        "complaints": len(pd.read_csv(DATA_DIR / "complaints.csv")),
        "app_events": len(pd.read_csv(DATA_DIR / "app_events.csv")),
        "incidents": len(pd.read_csv(DATA_DIR / "incidents.csv")),
        "drivers": len(pd.read_csv(DATA_DIR / "drivers.csv")),
        "vehicles": len(pd.read_csv(DATA_DIR / "vehicles.csv")),
        "hubs": len(pd.read_csv(DATA_DIR / "hubs.csv")),
    }
    descriptions = [
        ("orders", counts["orders"], "Customer demand creation, promised windows, route intent, and commercial value."),
        ("deliveries", counts["deliveries"], "Execution records linking drivers, vehicles, hubs, status outcomes, and route overrides."),
        ("customers", counts["customers"], "Customer profile, loyalty, engagement, and channel preference information."),
        ("complaints", counts["complaints"], "Post-service dissatisfaction data with complaint type, severity, and compensation."),
        ("app_events", counts["app_events"], "Digital interaction trail for latency, retry, tracking, and app-journey analysis."),
        ("incidents", counts["incidents"], "Operational exception records such as proof issues, battery alerts, and sync errors."),
        ("drivers", counts["drivers"], "Driver capability, ratings, and experience context."),
        ("vehicles", counts["vehicles"], "Fleet health, battery state, telematics version, and maintenance status."),
        ("hubs", counts["hubs"], "Hub location, capacity, and role in the operating network."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Source File", "Records", "Business Role"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "FDE68A")
    for name, records, role in descriptions:
        row = table.add_row().cells
        set_cell_text(row[0], name)
        set_cell_text(row[1], records)
        set_cell_text(row[2], role, font_size=10)


def add_quality_table(doc: Document) -> None:
    quality = read_csv_rows(OUTPUTS_DIR / "dataset_quality_summary.csv")
    labels = {
        "customers_preferred_channel_missing": "Missing preferred customer channel",
        "orders_booking_channel_missing": "Missing order booking channel",
        "deliveries_completed_timestamp_missing": "Missing delivery completion timestamp",
        "deliveries_rating_missing": "Missing post-delivery rating",
        "orders_without_delivery_record": "Orders without a delivery record",
    }
    interpretation = {
        "customers_preferred_channel_missing": "Weakens customer communication and segmentation analysis.",
        "orders_booking_channel_missing": "Reduces confidence in channel-attribution reporting.",
        "deliveries_completed_timestamp_missing": "Prevents accurate cycle-time measurement for part of the network.",
        "deliveries_rating_missing": "Removes part of the customer-experience feedback loop.",
        "orders_without_delivery_record": "Represents the most serious visibility failure in the dataset.",
    }
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Data Quality Issue", "Count", "Analytical Consequence"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "FDE68A")
    for row in quality:
        metric = row["metric"]
        cells = table.add_row().cells
        set_cell_text(cells[0], labels[metric])
        set_cell_text(cells[1], row["value"])
        set_cell_text(cells[2], interpretation[metric], font_size=10)


def add_image(doc: Document, path: Path, width: float, caption: str) -> None:
    doc.add_picture(str(path), width=Inches(width))
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(caption)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_intro(doc: Document) -> None:
    doc.add_paragraph("Chapter 1: Introduction and Assessment Context", style="Heading 1")
    doc.add_paragraph("1.1 Business Context", style="Heading 2")
    doc.add_paragraph(
        "NorthStar Urban Mobility and Logistics operates across several city zones and combines different service models including passenger journeys, parcel deliveries, retail movement, "
        "business fulfilment, and medical transport. The case study describes a business that has scaled quickly but has not developed a coherent reporting architecture to match that growth. "
        "As a result, service failures appear in multiple forms: late or failed deliveries, disconnected customer complaints, inconsistent digital experience, and fragmented operational oversight."
    )
    doc.add_paragraph("1.2 Assignment Focus and Learning Outcome Alignment", style="Heading 2")
    doc.add_paragraph(
        "The written coursework is designed to evidence the database and analytics skills required in the module guide. The SQL in R component addresses structured querying and analytical interpretation, "
        "the Python and MongoDB sections address document-oriented database thinking and integration with code, and the optimisation section addresses the ability to justify indexing and query design choices. "
        "The report is therefore organised to reflect the logic of the assessment rather than simply listing charts in isolation."
    )
    doc.add_paragraph("1.3 Aim", style="Heading 2")
    doc.add_paragraph(
        "The aim of this report is to diagnose NorthStar's most important operational problems, explain those problems using reproducible analysis, and propose a technically justified database redesign "
        "that improves both day-to-day reporting and future scalability."
    )
    doc.add_paragraph("1.4 Objectives", style="Heading 2")
    objectives = [
        "Use relational querying to identify reliability, visibility, and hub-performance failures.",
        "Clean and standardise the dataset so that inconsistent labels do not distort results.",
        "Engineer cross-file analytical features that reveal deeper operational risk patterns.",
        "Design a MongoDB-style case model that keeps complaints, app events, and delivery exceptions together.",
        "Propose an indexing and optimisation strategy grounded in realistic operational query paths.",
    ]
    for item in objectives:
        doc.add_paragraph(item, style="List Bullet")


def add_methodology(doc: Document) -> None:
    doc.add_paragraph("Chapter 2: Case Study Review, Data Inventory, and Methodology", style="Heading 1")
    doc.add_paragraph("2.1 Why the Case Demands an Integrated Approach", style="Heading 2")
    doc.add_paragraph(
        "The case study does not describe one single technical problem. Instead, it presents a business whose orders, deliveries, complaints, incidents, and app interactions are recorded separately and then "
        "interpreted through incomplete reporting. That means the analytical task is not only to produce summaries, but also to reconnect the operational story across those systems. This is why the report combines "
        "relational analysis with document-oriented redesign rather than assuming one database style can solve every requirement."
    )
    doc.add_paragraph("2.2 Dataset Inventory", style="Heading 2")
    add_dataset_inventory_table(doc)
    doc.add_paragraph(
        "The dataset contains 1,250 orders but only 950 delivery records, which immediately suggests that the operational pipeline is incomplete even before any advanced analysis is performed. "
        "The supporting files add the context needed to understand why this matters: complaints explain customer reaction, app events capture the digital journey, and incidents record operational exceptions."
    )
    doc.add_paragraph("2.3 Data Quality Audit", style="Heading 2")
    add_quality_table(doc)
    doc.add_paragraph(
        "The missing-data profile is analytically important because absence is not random. For example, missing proof-of-completion is strongly associated with poor outcomes, while missing delivery completion timestamps "
        "restrict accurate cycle-time analysis. The largest defect is the 300 orders with no delivery record, because this breaks traceability between commercial demand and operational execution."
    )
    doc.add_paragraph("2.4 Cleaning and Integration Method", style="Heading 2")
    doc.add_paragraph(
        "The first preprocessing step was standardisation of location labels. Without that step, values such as Central, CENTRAL, and Ctr would be treated as separate categories, producing misleading zone-level summaries. "
        "The second step was cross-file integration so that delivery status could be read together with order value, hub identity, compensation, vehicle condition, and app-event context. The third step was feature engineering, "
        "including completion hours, net value after compensation, route-override bands, and proof-of-completion flags."
    )
    add_image(
        doc,
        EVIDENCE_DIR / "python_merged_preview.png",
        6.35,
        "Figure 1. Python merged dataset preview used to bring delivery, order, and hub information into one analytical frame.",
    )
    doc.add_paragraph(
        "Figure 1 shows the integrated working dataset that underpins the later analysis. This step is methodologically important because it changes the problem from a simple file-by-file review into a joined operational dataset "
        "in which business performance can be traced end to end."
    )
    doc.add_page_break()


def add_sql_r_section(doc: Document) -> None:
    doc.add_paragraph("Chapter 3: SQL in R Analysis and Relational Findings", style="Heading 1")
    doc.add_paragraph("3.1 Why SQL in R Was Appropriate", style="Heading 2")
    doc.add_paragraph(
        "The first analytical layer uses SQL in an R-oriented environment because the core operational records are structured and benefit from joins, filters, groupings, and direct quantitative summaries. "
        "This method is especially suitable for the module requirement because it demonstrates that SQL can be embedded within a wider analytical workflow rather than treated as a standalone querying exercise."
    )
    doc.add_paragraph("3.2 Hub Reliability Risk", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "sql_hub_risk_query.png",
        6.45,
        "Figure 2. SQL query result ranking hubs by non-on-time delivery share, failed-delivery share, route overrides, and completion hours.",
    )
    doc.add_paragraph(
        "Figure 2 shows that Central Core has the worst overall reliability profile, with 41.7% of its deliveries arriving either delayed or failed, while Midtown Relay also performs poorly with a 37.5% non-on-time rate. "
        "The importance of this result is not only that some hubs perform badly, but that the weakest hubs are clustered in the central operating layer. This suggests a structural capacity or coordination problem rather than random noise."
    )
    doc.add_paragraph(
        "The same query also shows that poor reliability is linked with longer completion hours and elevated route overrides, which indicates operational friction in the dispatch or routing process. East Dock is the most stable hub in the table, "
        "providing a useful internal benchmark for what more controlled performance looks like."
    )
    doc.add_paragraph("3.3 Orders Missing a Delivery Record", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "sql_missing_orders_query.png",
        6.35,
        "Figure 3. SQL query result for orders that do not have a linked delivery record, broken down by service type.",
    )
    doc.add_paragraph(
        "Figure 3 confirms that the visibility problem is not confined to one narrow service category. Passenger, Parcel, and Retail services all contribute heavily to the missing-order gap, with 79, 78, and 73 missing records respectively. "
        "This matters because it changes the management response. If the issue were concentrated in one service type, the business could investigate one specialist workflow. Instead, the evidence suggests a broader system integration failure."
    )
    doc.add_paragraph("3.4 Zone Performance After Cleaning Inconsistent Labels", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "sql_zone_performance_table.png",
        6.35,
        "Figure 4. Zone-level performance table after cleaning inconsistent location labels across source files.",
    )
    add_image(
        doc,
        EVIDENCE_DIR / "sql_zone_failed_chart.png",
        6.1,
        "Figure 5. Failed delivery share by operating zone, showing Central as the weakest service layer.",
    )
    doc.add_paragraph(
        "Once inconsistent zone labels were standardised, the Central zone emerged as the weakest performing area in the network. It records only 60.5% on-time delivery and 20.2% failed delivery, which is materially worse than East, North, or South. "
        "This demonstrates the value of cleaning before interpretation: without standardisation, the Central pattern would be diluted across multiple labels and the most important finding in the network would be understated."
    )
    doc.add_paragraph("3.5 Customer Dissatisfaction Is Not Explained by Failure Alone", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "sql_complaint_mix_chart.png",
        6.25,
        "Figure 6. Complaint mix by delivery status, showing that dissatisfaction also appears against orders marked OnTime.",
    )
    doc.add_paragraph(
        "Figure 6 is a particularly important relational insight because it shows that the customer-experience problem is wider than simple failed deliveries. Delay complaints are common even on deliveries recorded as OnTime, and AppIssue, DriverBehaviour, "
        "and MissedPickup complaints also appear in that category. This suggests that NorthStar's internal service-closure logic does not fully match the customer's lived experience."
    )
    doc.add_paragraph(
        "Taken together, the SQL in R analysis demonstrates that the structured relational layer is already capable of exposing serious operational weaknesses when used carefully. It also provides a strong foundation for the later Python and MongoDB stages, "
        "because those stages are responding to real patterns rather than inventing a redesign without evidence."
    )
    doc.add_page_break()


def add_python_section(doc: Document) -> None:
    doc.add_paragraph("Chapter 4: Python Analytics, Feature Engineering, and Interpretation", style="Heading 1")
    doc.add_paragraph("4.1 Why Python Was Added to the Workflow", style="Heading 2")
    doc.add_paragraph(
        "Python was used after the SQL stage because several of the most useful features in the NorthStar dataset are engineered rather than directly stored. Python is well suited to this stage because it allows flexible joining, cleaning, conditional logic, "
        "and rapid creation of interpretable summary frames. This extends the analytical value of the dataset beyond simple grouped counts."
    )
    doc.add_paragraph("4.2 Proof of Completion as a Risk Signal", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "python_proof_risk_table.png",
        6.25,
        "Figure 7. Proof-of-completion risk comparison showing sharply worse outcomes when proof is missing.",
    )
    doc.add_paragraph(
        "Figure 7 is one of the strongest findings in the whole report. Deliveries with proof of completion missing show a 65.2% delayed rate and a 34.8% failed rate, compared with 17.8% and 12.3% when proof is present. The average customer rating also falls "
        "from 3.92 to 3.17. This means proof of completion is not simply an administrative afterthought. In this dataset it behaves like an operational risk flag and should be treated as such."
    )
    doc.add_paragraph("4.3 Maintenance State and Fleet-Control Weakness", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "python_maintenance_table.png",
        6.25,
        "Figure 8. Vehicle maintenance performance summary showing the penalty attached to operating InRepair vehicles.",
    )
    doc.add_paragraph(
        "Figure 8 shows that deliveries linked to vehicles marked InRepair achieve only 49.2% on-time completion and 30.3% failure. That is dramatically worse than the Active and Scheduled categories. The management implication is that maintenance status is not "
        "being integrated tightly enough into operational dispatch rules. If a vehicle already classified as InRepair continues to carry live workload, NorthStar is effectively converting a maintenance issue into a customer-experience issue."
    )
    doc.add_paragraph("4.4 Complaint Intensity by Hub", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "python_complaints_by_hub_chart.png",
        6.2,
        "Figure 9. Complaint intensity by hub measured as complaints per delivery.",
    )
    doc.add_paragraph(
        "Figure 9 adds a customer-facing layer to the earlier hub analysis. Riverside Hub, East Dock, Midtown Relay, and Central Core all show elevated complaint intensity when complaints are normalised by delivery volume. This is useful because raw complaint counts "
        "alone can overstate high-volume hubs and understate smaller but weaker ones. Complaint intensity offers a fairer view of customer pain."
    )
    doc.add_paragraph("4.5 Digital Friction and App Latency", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "python_latency_chart.png",
        6.25,
        "Figure 10. Highest app latency segments by zone context and event type.",
    )
    doc.add_paragraph(
        "The case study is not purely about vehicles and hubs. Figure 10 shows that digital friction is also significant, especially in Airport and Central journeys. Track-order, chat-opened, and payment-retry events appear among the highest latency segments, "
        "which helps explain why AppIssue complaints occur even when some deliveries are technically marked OnTime. For NorthStar, service quality is partly operational and partly digital."
    )
    doc.add_paragraph("4.6 Profitability and the Risk of Misleading Success Signals", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "python_profitability_chart.png",
        6.25,
        "Figure 11. Average net value after direct cost and compensation by hub.",
    )
    doc.add_paragraph(
        "Figure 11 shows an important managerial tension. Central Core still appears commercially strong on a simple net-value basis despite recording the worst reliability profile in the earlier SQL analysis. This illustrates why purely financial summaries can be dangerous: "
        "a hub may remain profitable in the short term while simultaneously damaging customer experience, increasing exceptions, and weakening long-term trust."
    )
    doc.add_paragraph(
        "Overall, the Python stage deepens the relational analysis by transforming operational fields into decision-relevant indicators. It shows that the most useful insights do not come only from counting statuses, but from combining statuses with proof flags, complaints, "
        "vehicle state, and digital events."
    )
    doc.add_page_break()


def add_collection_table(doc: Document) -> None:
    rows = [
        ("customer_cases", "Customer profile, order header, delivery summary, complaints, app events, and current case status.", "Keeps the complete operational story for one case together."),
        ("service_orders", "Structured order facts such as service type, route intent, promised window, and commercial values.", "Retains a clean document representation of the commercial demand layer."),
        ("delivery_events", "Dispatch events, incident snapshots, route overrides, proof-of-completion state, and operational milestones.", "Supports timeline review and operational investigation."),
        ("asset_health", "Vehicle metadata, maintenance history, battery health, and recent linked service outcomes.", "Separates long-lived fleet state from individual customer cases."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Collection", "Document Focus", "Reason for Inclusion"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "FDE68A")
    for name, focus, reason in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], name)
        set_cell_text(cells[1], focus, font_size=9.8)
        set_cell_text(cells[2], reason, font_size=9.8)


def add_index_table(doc: Document) -> None:
    rows = [
        ("customer_cases(order_id)", "Single-case order lookup", "Fast retrieval of one case from an operational order reference."),
        ("customer_cases(customer.customer_id, case_status)", "Customer case history review", "Supports open and closed case filtering per customer."),
        ("customer_cases(delivery_summary.hub_id, delivery_summary.delivery_status)", "Hub-level exception triage", "Supports operational dashboards for delayed or failed cases by hub."),
        ("delivery_events(vehicle_id, maintenance_status)", "Vehicle reliability analysis", "Links delivery exceptions to asset condition."),
        ("delivery_events(delivery_summary.proof_of_completion_missing)", "Proof-risk monitoring", "Supports exception rules around missing proof records."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Index", "Primary Query Path", "Optimisation Benefit"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "FDE68A")
    for name, path, benefit in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], name, font_size=9.6)
        set_cell_text(cells[1], path, font_size=9.8)
        set_cell_text(cells[2], benefit, font_size=9.8)


def add_mongo_section(doc: Document) -> None:
    doc.add_paragraph("Chapter 5: MongoDB Design, Indexing Strategy, and Query Optimisation", style="Heading 1")
    doc.add_paragraph("5.1 Why a Document Model Is Needed", style="Heading 2")
    doc.add_paragraph(
        "The NorthStar case contains several domains that do not fit comfortably into a rigid relational structure when the objective is day-to-day operational investigation. Complaint journeys, app-event trails, incident chains, and evolving customer cases all have a nested "
        "or time-sequenced character. Splitting those histories across many tables can preserve normalisation, but it makes practical investigation harder. A document model is better for keeping related evidence together."
    )
    doc.add_paragraph("5.2 Proposed Collection Structure", style="Heading 2")
    add_collection_table(doc)
    doc.add_paragraph(
        "The proposed structure keeps core customer-case evidence in a `customer_cases` collection while allowing other collections to support operational order history and long-lived asset state. This is a more realistic fit for NorthStar than trying to force every evolving story "
        "into separate relational entities."
    )
    doc.add_paragraph("5.3 Example Nested Case Document", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "mongo_case_document.png",
        6.05,
        "Figure 12. Proposed MongoDB customer case document combining service order, delivery outcome, complaints, and app events.",
    )
    doc.add_paragraph(
        "Figure 12 shows the logic of the document model. The case can be understood in one place: the customer context, the commercial order, the delivery outcome, the complaint record, and the digital interaction trail are all visible together. "
        "This is particularly valuable for support teams and operational analysts, because it reduces the need for repeated multi-table reconstruction of the same incident."
    )
    doc.add_paragraph("5.4 Index Design", style="Heading 2")
    add_index_table(doc)
    add_image(
        doc,
        EVIDENCE_DIR / "mongo_index_design.png",
        6.05,
        "Figure 13. MongoDB index design evidence taken from the Python and MongoDB workflow.",
    )
    doc.add_paragraph(
        "The indexing strategy is deliberately driven by business access paths rather than by generic database advice. NorthStar is likely to ask questions such as: Which open cases belong to this customer? Which failed cases are accumulating in Central Core? Which vehicles with "
        "maintenance issues are also associated with delivery exceptions? The proposed indexes map directly to those use cases."
    )
    doc.add_paragraph("5.5 Query Optimisation and Aggregation Logic", style="Heading 2")
    add_image(
        doc,
        EVIDENCE_DIR / "mongo_pipeline_results.png",
        6.15,
        "Figure 14. MongoDB aggregation output for open delayed or failed Central Core cases.",
    )
    doc.add_paragraph(
        "Figure 14 demonstrates the type of document-oriented triage query that would be far more awkward in a fragmented relational design. The aggregation filters by hub, delivery outcome, and case status, then returns the most operationally useful summary fields for intervention. "
        "In practice, this enables support and operations teams to identify problematic cases quickly without reconstructing the same story from multiple unrelated systems."
    )
    doc.add_paragraph("5.6 Explain-Plan Discussion", style="Heading 2")
    doc.add_paragraph(
        "In a live MongoDB Atlas environment, the next optimisation step would be to capture `explain('executionStats')` before and after index creation. The purpose would be to demonstrate reduced collection scans, fewer examined documents, and improved execution time for the main case-triage queries. "
        "In this workspace the notebook was executed using a MongoDB-compatible fallback client because no live Atlas cluster connection was available. The commands and index logic are therefore evidenced, but the final Atlas execution statistics still require a live cluster if the tutor specifically expects them."
    )
    doc.add_page_break()


def add_recommendations(doc: Document) -> None:
    doc.add_paragraph("Chapter 6: Recommendations and Conclusion", style="Heading 1")
    doc.add_paragraph("6.1 Priority Recommendations", style="Heading 2")
    recommendations = [
        "Treat the 300 orders without a delivery record as a high-priority control breach and investigate the workflow stage at which traceability is being lost.",
        "Focus operational improvement first on Central Core and Midtown Relay, because the central network layer combines high non-on-time performance with high failure rates.",
        "Promote missing proof of completion from a passive data field to an active exception trigger in operational dashboards.",
        "Block or justify dispatch allocation to vehicles already marked InRepair, because the performance penalty is too large to ignore.",
        "Integrate app latency monitoring into service management, especially for Airport and Central journeys where digital friction is more visible.",
        "Adopt a hybrid architecture in which structured operational facts remain relational but semi-structured case evidence is stored in MongoDB documents.",
    ]
    for item in recommendations:
        doc.add_paragraph(item, style="List Number")
    doc.add_paragraph("6.2 Conclusion", style="Heading 2")
    doc.add_paragraph(
        "The NorthStar case is best understood as a visibility and integration problem rather than a simple late-delivery problem. The relational layer already reveals serious weaknesses, including the scale of missing delivery records and the central hub network's poor reliability. "
        "Python deepens the diagnosis by showing how risk intensifies when proof is missing, when vehicle maintenance state is weak, and when customer journeys suffer digital friction. MongoDB then becomes relevant not as a replacement for all relational storage, but as the right design for nested case histories, "
        "complaint trails, and event chains that need to stay together. The strongest final recommendation is therefore a hybrid model: keep relational reporting for structured facts, strengthen analytics through cleaned and engineered Python workflows, and redesign semi-structured operational histories into indexed "
        "MongoDB documents for faster, more realistic case management."
    )


def add_references(doc: Document) -> None:
    doc.add_page_break()
    doc.add_paragraph("References", style="Heading 1")
    refs = [
        "University of West London. Databases and Analytics Module Study Guide, Academic Year 2025-2026.",
        "University of West London. Databases and Analytics Assignment Brief, Semester 2, 2025-2026.",
        "NorthStar Urban Mobility and Logistics Case Study, Databases and Analytics Module, Semester 2, 2026.",
        "NorthStar_dataset: orders, deliveries, customers, complaints, app_events, incidents, drivers, vehicles, hubs, and data dictionary files supplied with the coursework.",
        "MongoDB Documentation. Query Optimisation, Indexing, and Aggregation Framework guidance.",
        "RSQLite, DBI, pandas, and PyMongo project documentation used to structure the executed analytical workflow.",
    ]
    for item in refs:
        doc.add_paragraph(item, style="List Bullet")


def add_appendices(doc: Document) -> None:
    doc.add_page_break()
    doc.add_paragraph("Appendix A: Evidence Figure Register", style="Heading 1")
    figure_rows = [
        ("Figure 1", "Python merged dataset preview", "Shows the integrated analytical base table before feature engineering."),
        ("Figure 2", "Hub reliability SQL query", "Ranks hubs by non-on-time share and failed-delivery share."),
        ("Figure 3", "Missing delivery link SQL query", "Shows the 300-order visibility gap distributed across service types."),
        ("Figure 4", "Zone performance table", "Summarises cleaned zone-level outcome rates."),
        ("Figure 5", "Failed delivery by zone chart", "Highlights Central as the weakest zone."),
        ("Figure 6", "Complaint mix chart", "Demonstrates dissatisfaction beyond failed deliveries alone."),
        ("Figure 7", "Proof-of-completion risk table", "Quantifies the performance penalty when proof is missing."),
        ("Figure 8", "Maintenance performance table", "Shows the operational cost of using InRepair vehicles."),
        ("Figure 9", "Complaint intensity by hub chart", "Normalises complaints by delivery volume."),
        ("Figure 10", "App latency chart", "Shows the digital-service dimension of the customer experience problem."),
        ("Figure 11", "Profitability chart", "Demonstrates why financial success alone can hide operational weakness."),
        ("Figure 12", "MongoDB case document", "Illustrates the proposed nested document structure."),
        ("Figure 13", "MongoDB index design", "Shows the index set aligned to real case-investigation queries."),
        ("Figure 14", "MongoDB aggregation output", "Demonstrates delayed or failed open-case triage."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Figure", "Title", "Purpose in Report"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
        shade_cell(table.rows[0].cells[idx], "FDE68A")
    for label, title, purpose in figure_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], label)
        set_cell_text(cells[1], title, font_size=10)
        set_cell_text(cells[2], purpose, font_size=10)

    doc.add_paragraph()
    doc.add_paragraph("Appendix B: Output Files Generated", style="Heading 1")
    outputs = [
        "Executed SQL in R notebook: notebooks/01_northstar_sql_r_analytics.ipynb",
        "Executed Python and MongoDB notebook: notebooks/02_northstar_python_mongodb_solution.ipynb",
        "Notebook HTML exports: notebooks/html/01_northstar_sql_r_analytics.html and notebooks/html/02_northstar_python_mongodb_solution.html",
        "Structured analytical outputs: artifacts/outputs/*.csv",
        "Evidence figures used in the report: report/assets/evidence/*.png",
        "Main coursework report deliverable: report/northstar_assignment_report.docx",
    ]
    for item in outputs:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    doc = Document()
    style_document(doc)
    add_footer(doc.sections[0])
    add_cover_page(doc)
    add_footer(doc.sections[0])
    add_front_matter(doc)
    add_abstract(doc)
    add_intro(doc)
    add_methodology(doc)
    add_sql_r_section(doc)
    add_python_section(doc)
    add_mongo_section(doc)
    add_recommendations(doc)
    add_references(doc)
    add_appendices(doc)
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
